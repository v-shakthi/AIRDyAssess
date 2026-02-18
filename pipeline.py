"""
ingestion/pipeline.py
=====================
Multi-format document ingestion pipeline.

Supports: PDF, DOCX, TXT
Flow: Extract text → Clean → Chunk → Embed → Store in ChromaDB

In production: extend with SharePoint connector, Confluence API, S3 bucket reader.
"""

import re
import hashlib
import logging
from pathlib import Path
from dataclasses import dataclass

import chromadb
from chromadb.utils import embedding_functions

from config import settings

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Extracted document
# ---------------------------------------------------------------------------

@dataclass
class ExtractedDocument:
    filename: str
    content: str
    page_count: int
    file_type: str


# ---------------------------------------------------------------------------
# Extractors
# ---------------------------------------------------------------------------

def extract_pdf(path: Path) -> ExtractedDocument:
    """Extract text from PDF using pypdf."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        pages = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                pages.append(f"[Page {i+1}]\n{text}")
        return ExtractedDocument(
            filename=path.name,
            content="\n\n".join(pages),
            page_count=len(reader.pages),
            file_type="pdf",
        )
    except Exception as e:
        raise ValueError(f"Failed to extract PDF {path.name}: {e}")


def extract_docx(path: Path) -> ExtractedDocument:
    """Extract text from DOCX preserving heading structure."""
    try:
        from docx import Document
        doc = Document(str(path))
        sections = []
        for para in doc.paragraphs:
            if para.text.strip():
                prefix = "## " if para.style.name.startswith("Heading") else ""
                sections.append(f"{prefix}{para.text}")
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    sections.append(row_text)
        return ExtractedDocument(
            filename=path.name,
            content="\n".join(sections),
            page_count=len(doc.paragraphs) // 25 + 1,  # Estimate
            file_type="docx",
        )
    except Exception as e:
        raise ValueError(f"Failed to extract DOCX {path.name}: {e}")


def extract_txt(path: Path) -> ExtractedDocument:
    """Extract plain text."""
    content = path.read_text(encoding="utf-8", errors="ignore")
    return ExtractedDocument(
        filename=path.name,
        content=content,
        page_count=max(1, len(content) // 3000),
        file_type="txt",
    )


def extract_document(path: Path) -> ExtractedDocument:
    """Route to the correct extractor based on file extension."""
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf(path)
    elif suffix in (".docx", ".doc"):
        return extract_docx(path)
    elif suffix in (".txt", ".md"):
        return extract_txt(path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}. Supported: pdf, docx, txt, md")


# ---------------------------------------------------------------------------
# Chunker
# ---------------------------------------------------------------------------

def chunk_text(
    text: str,
    chunk_size: int = None,
    chunk_overlap: int = None,
) -> list[str]:
    """
    Splits text into overlapping chunks.
    Tries to split on paragraph boundaries first, then sentence boundaries.
    """
    chunk_size = chunk_size or settings.chunk_size
    chunk_overlap = chunk_overlap or settings.chunk_overlap

    # Split on double newlines (paragraphs) first
    paragraphs = [p.strip() for p in re.split(r'\n\n+', text) if p.strip()]

    chunks = []
    current = ""

    for para in paragraphs:
        if len(current) + len(para) <= chunk_size:
            current = f"{current}\n\n{para}".strip()
        else:
            if current:
                chunks.append(current)
            # If the paragraph itself is too long, split it further
            if len(para) > chunk_size:
                words = para.split()
                sub = ""
                for word in words:
                    if len(sub) + len(word) + 1 <= chunk_size:
                        sub = f"{sub} {word}".strip()
                    else:
                        if sub:
                            chunks.append(sub)
                        sub = word
                if sub:
                    current = sub
            else:
                current = para

    if current:
        chunks.append(current)

    # Add overlap: prepend tail of previous chunk to next
    overlapped = []
    for i, chunk in enumerate(chunks):
        if i == 0:
            overlapped.append(chunk)
        else:
            tail = chunks[i - 1][-chunk_overlap:] if len(chunks[i - 1]) > chunk_overlap else chunks[i - 1]
            overlapped.append(f"{tail}\n{chunk}")

    return overlapped


# ---------------------------------------------------------------------------
# ChromaDB store
# ---------------------------------------------------------------------------

def get_collection(session_id: str):
    """Get or create a ChromaDB collection scoped to a session."""
    client = chromadb.PersistentClient(path=settings.chroma_persist_dir)
    ef = embedding_functions.DefaultEmbeddingFunction()
    collection_name = f"{settings.collection_name}_{session_id[:8]}"
    return client.get_or_create_collection(
        name=collection_name,
        embedding_function=ef,
    )


def ingest_documents(
    file_paths: list[Path],
    session_id: str,
    progress_callback=None,
) -> tuple[list[ExtractedDocument], int]:
    """
    Full ingestion pipeline: extract → chunk → embed → store.

    Args:
        file_paths: List of file paths to ingest
        session_id: Unique session ID for ChromaDB scoping
        progress_callback: Optional callable(step: str, pct: int)

    Returns:
        (list of ExtractedDocuments, total chunk count)
    """
    collection = get_collection(session_id)
    extracted_docs = []
    total_chunks = 0

    for i, path in enumerate(file_paths):
        if progress_callback:
            pct = int((i / len(file_paths)) * 40)
            progress_callback(f"Extracting {path.name}...", pct)

        try:
            doc = extract_document(path)
            extracted_docs.append(doc)
            logger.info(f"Extracted {path.name}: {len(doc.content)} chars, {doc.page_count} pages")
        except Exception as e:
            logger.error(f"Skipping {path.name}: {e}")
            continue

        # Chunk
        chunks = chunk_text(doc.content)
        logger.info(f"Chunked {path.name} into {len(chunks)} chunks")

        # Store in ChromaDB
        ids, documents, metadatas = [], [], []
        for j, chunk in enumerate(chunks):
            chunk_id = hashlib.md5(f"{session_id}:{path.name}:{j}".encode()).hexdigest()
            ids.append(chunk_id)
            documents.append(chunk)
            metadatas.append({
                "source": path.name,
                "file_type": doc.file_type,
                "chunk_index": j,
                "session_id": session_id,
            })

        # Upsert in batches of 100
        batch_size = 100
        for b in range(0, len(ids), batch_size):
            collection.upsert(
                ids=ids[b:b+batch_size],
                documents=documents[b:b+batch_size],
                metadatas=metadatas[b:b+batch_size],
            )

        total_chunks += len(chunks)

    if progress_callback:
        progress_callback("Documents indexed in vector store.", 45)

    return extracted_docs, total_chunks


def search_documents(query: str, session_id: str, n_results: int = None) -> list[dict]:
    """
    Semantic search over ingested documents for a given session.
    Returns list of {content, source, distance} dicts.
    """
    n_results = n_results or settings.top_k_retrieval
    try:
        collection = get_collection(session_id)
        results = collection.query(query_texts=[query], n_results=min(n_results, collection.count()))
        out = []
        for i, doc in enumerate(results["documents"][0]):
            out.append({
                "content": doc,
                "source": results["metadatas"][0][i].get("source", "unknown"),
                "distance": results["distances"][0][i],
            })
        return out
    except Exception as e:
        logger.error(f"Search error: {e}")
        return []
